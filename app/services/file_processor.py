import io
from typing import Union, Dict, Any
import fitz  # PyMuPDF
from docx import Document
from PIL import Image

from app.core.exceptions import FileProcessingError
from app.utils.file_helpers import validate_file_type, validate_file_size


class FileProcessor:
    """Handle processing of different file types"""
    
    async def process_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Process uploaded file based on its type
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            
        Returns:
            Dict containing processed content and metadata
        """
        # Validate file
        validate_file_size(len(file_content))
        file_type = validate_file_type(filename)
        
        try:
            if file_type == "text":
                return await self._process_text_file(file_content, filename)
            elif file_type == "document":
                return await self._process_document_file(file_content, filename)
            elif file_type == "image":
                return await self._process_image_file(file_content, filename)
        except Exception as e:
            raise FileProcessingError(f"Failed to process file '{filename}': {str(e)}")
    
    async def _process_text_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process text files (.txt, .md)"""
        try:
            text_content = file_content.decode('utf-8')
            return {
                "type": "text",
                "content": text_content,
                "filename": filename
            }
        except UnicodeDecodeError:
            raise FileProcessingError(f"Unable to decode text file '{filename}' as UTF-8")
    
    async def _process_document_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process document files (.pdf, .docx)"""
        if filename.lower().endswith('.pdf'):
            return await self._process_pdf(file_content, filename)
        elif filename.lower().endswith('.docx'):
            return await self._process_docx(file_content, filename)
    
    async def _process_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from PDF using PyMuPDF"""
        try:
            pdf_document = fitz.open(stream=file_content, filetype="pdf")
            text_content = ""
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                text_content += page.get_text()
            
            pdf_document.close()
            
            if not text_content.strip():
                raise FileProcessingError("No text content found in PDF")
            
            return {
                "type": "text",
                "content": text_content,
                "filename": filename,
                "pages": pdf_document.page_count if 'pdf_document' in locals() else 0
            }
        except Exception as e:
            raise FileProcessingError(f"Failed to process PDF '{filename}': {str(e)}")
    
    async def _process_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from Word document"""
        try:
            doc = Document(io.BytesIO(file_content))
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            if not text_content.strip():
                raise FileProcessingError("No text content found in Word document")
            
            return {
                "type": "text",
                "content": text_content,
                "filename": filename,
                "paragraphs": len(doc.paragraphs)
            }
        except Exception as e:
            raise FileProcessingError(f"Failed to process Word document '{filename}': {str(e)}")
    
    async def _process_image_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process image files for multimodal AI processing"""
        try:
            # Validate image can be opened
            image = Image.open(io.BytesIO(file_content))
            width, height = image.size
            format_name = image.format
            
            return {
                "type": "image",
                "content": file_content,  # Keep raw bytes for Gemini
                "filename": filename,
                "width": width,
                "height": height,
                "format": format_name
            }
        except Exception as e:
            raise FileProcessingError(f"Failed to process image '{filename}': {str(e)}")